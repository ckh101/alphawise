"""
文件解析工具
支持 PDF、Word、Excel、Markdown、TXT 格式
"""

import io
from pathlib import Path

from harness.core.logger import get_logger

logger = get_logger(__name__)

MAX_CHARS = 50000
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".md", ".txt"}


def parse_file(filename: str, content_bytes: bytes) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {ext}")

    if ext == ".pdf":
        text = _parse_pdf(content_bytes)
    elif ext in (".docx", ".doc"):
        text = _parse_docx(content_bytes)
    elif ext in (".xlsx", ".xls"):
        text = _parse_xlsx(content_bytes)
    else:
        text = content_bytes.decode("utf-8", errors="replace")

    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[内容过长，已截断]"
    return text


def _parse_pdf(data: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                parts.append(text)
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    parts.append(" | ".join(str(c or "") for c in row))
    return "\n\n".join(parts)


def _parse_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _parse_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            parts.append(" | ".join(str(c or "") for c in row))
    wb.close()
    return "\n".join(parts)
